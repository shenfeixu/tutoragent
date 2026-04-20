import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def export_markdown_to_docx(md_content: str, title: str = "商业计划书", subtitle: str = "") -> bytes:
    """
    将 Markdown 字符串转换为 Word 文档字节流。
    支持：标题 (H1, H2, H3)、加粗、列表、正文。
    """
    doc = Document()
    
    # 1. 添加封面/大标题
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if subtitle:
        sub = doc.add_paragraph(subtitle)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.runs[0]
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_page_break()

    # 2. 解析正文
    lines = md_content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # 标题解析
        if line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=2)
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=1)
        elif line.startswith('# '):
            doc.add_heading(line.replace('# ', ''), level=0)
            
        # 列表解析
        elif line.startswith('* ') or line.startswith('- '):
            clean_text = line[2:]
            p = doc.add_paragraph(style='List Bullet')
            _handle_formatting(p, clean_text)
        elif re.match(r'^\d+\.', line):
            clean_text = re.sub(r'^\d+\.\s*', '', line)
            p = doc.add_paragraph(style='List Number')
            _handle_formatting(p, clean_text)
            
        # 普通段落
        else:
            p = doc.add_paragraph()
            _handle_formatting(p, line)

    # 保存到字节流
    target = io.BytesIO()
    doc.save(target)
    return target.getvalue()

def _handle_formatting(paragraph, text):
    """
    处理基本的加粗格式 **text**
    """
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

